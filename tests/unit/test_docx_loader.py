"""Unit tests for DOCX Loader contract and behavior.

Tests verify:
- DocxLoader initialization and load() signature
- Document structure and metadata (source_path, doc_type, doc_hash, paragraph_count)
- Heading and paragraph conversion to Markdown
- Inline image extraction, placeholder insertion, and metadata recording
- Error handling for invalid inputs (wrong extension, missing file)
"""

from pathlib import Path

import pytest

from src.core.types import Document
from src.libs.loader.base_loader import BaseLoader
from src.libs.loader.docx_loader import DocxLoader


def _create_test_docx(tmp_path: Path, filename: str = "test.docx") -> Path:
    """Create a minimal DOCX file with headings and paragraphs using python-docx."""
    from docx import Document as DocxDocument

    path = tmp_path / filename
    doc = DocxDocument()
    doc.add_paragraph("Title Paragraph")
    doc.add_heading("Main Heading", level=1)
    doc.add_paragraph("Body text under heading.")
    doc.add_heading("Sub Heading", level=2)
    doc.add_paragraph("More content here.")
    doc.save(str(path))
    return path


def _create_docx_with_image(tmp_path: Path, filename: str = "with_image.docx") -> Path:
    """Create a DOCX file containing a 1x1 red PNG inline image."""
    import struct
    import zlib

    from docx import Document as DocxDocument
    from docx.shared import Inches

    # Build a minimal 1x1 red PNG in memory
    def _png_chunk(chunk_type: bytes, data: bytes) -> bytes:
        crc = zlib.crc32(chunk_type + data) & 0xFFFFFFFF
        return struct.pack(">I", len(data)) + chunk_type + data + struct.pack(">I", crc)

    png_bytes = (
        b"\x89PNG\r\n\x1a\n"
        + _png_chunk(b"IHDR", struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0))
        + _png_chunk(b"IDAT", zlib.compress(b"\x00\xff\x00\x00"))
        + _png_chunk(b"IEND", b"")
    )

    png_path = tmp_path / "test_image.png"
    png_path.write_bytes(png_bytes)

    path = tmp_path / filename
    doc = DocxDocument()
    doc.add_paragraph("Paragraph before image.")
    doc.add_picture(str(png_path), width=Inches(1))
    doc.add_paragraph("Paragraph after image.")
    doc.save(str(path))
    return path


class TestDocxLoaderInitialization:
    """Tests for DocxLoader initialization."""

    def test_can_instantiate(self):
        """DocxLoader can be initialized when python-docx is available."""
        loader = DocxLoader()
        assert loader is not None

    def test_inherits_from_base_loader(self):
        """DocxLoader inherits from BaseLoader."""
        loader = DocxLoader()
        assert isinstance(loader, BaseLoader)

    def test_default_extract_images_is_true(self):
        """extract_images defaults to True."""
        loader = DocxLoader()
        assert loader.extract_images is True

    def test_extract_images_can_be_disabled(self):
        """extract_images=False disables image extraction."""
        loader = DocxLoader(extract_images=False)
        assert loader.extract_images is False

    def test_custom_image_storage_dir(self, tmp_path):
        """image_storage_dir is stored as Path."""
        loader = DocxLoader(image_storage_dir=str(tmp_path))
        assert loader.image_storage_dir == tmp_path


class TestDocxLoaderValidation:
    """Tests for input validation."""

    def test_load_requires_docx_extension(self, tmp_path):
        """load() raises ValueError for non-DOCX files."""
        txt_file = tmp_path / "test.txt"
        txt_file.write_text("not a docx")

        loader = DocxLoader()
        with pytest.raises(ValueError, match="not a DOCX"):
            loader.load(txt_file)

    def test_load_requires_docx_extension_pdf(self, tmp_path):
        """load() raises ValueError for PDF files."""
        pdf_file = tmp_path / "test.pdf"
        pdf_file.write_bytes(b"%PDF-1.4 dummy")

        loader = DocxLoader()
        with pytest.raises(ValueError, match="not a DOCX"):
            loader.load(pdf_file)

    def test_load_nonexistent_file(self):
        """load() raises FileNotFoundError for missing files."""
        loader = DocxLoader()
        with pytest.raises(FileNotFoundError):
            loader.load("nonexistent.docx")


class TestDocxLoaderConversion:
    """Tests for DOCX to Markdown conversion."""

    def test_load_returns_document_with_metadata(self, tmp_path):
        """load() returns Document with correct metadata."""
        docx_path = _create_test_docx(tmp_path)

        loader = DocxLoader()
        doc = loader.load(docx_path)

        assert isinstance(doc, Document)
        assert doc.id.startswith("doc_")
        assert doc.metadata["source_path"] == str(docx_path)
        assert doc.metadata["doc_type"] == "docx"
        assert "doc_hash" in doc.metadata
        assert doc.metadata["paragraph_count"] >= 4

    def test_headings_converted_to_markdown(self, tmp_path):
        """Headings are converted to Markdown # ## ### format."""
        docx_path = _create_test_docx(tmp_path)

        loader = DocxLoader()
        doc = loader.load(docx_path)

        assert "# Main Heading" in doc.text
        assert "## Sub Heading" in doc.text

    def test_paragraphs_preserved(self, tmp_path):
        """Paragraph text is extracted and preserved."""
        docx_path = _create_test_docx(tmp_path)

        loader = DocxLoader()
        doc = loader.load(docx_path)

        assert "Title Paragraph" in doc.text
        assert "Body text under heading." in doc.text
        assert "More content here." in doc.text

    def test_document_hash_consistency(self, tmp_path):
        """Same DOCX produces same document hash and id."""
        docx_path = _create_test_docx(tmp_path)

        loader = DocxLoader()
        doc1 = loader.load(docx_path)
        doc2 = loader.load(docx_path)

        assert doc1.metadata["doc_hash"] == doc2.metadata["doc_hash"]
        assert doc1.id == doc2.id

    def test_extract_title_from_heading(self, tmp_path):
        """Title is extracted from first heading when present."""
        docx_path = _create_test_docx(tmp_path)

        loader = DocxLoader()
        doc = loader.load(docx_path)

        assert "title" in doc.metadata
        assert doc.metadata["title"] is not None


class TestDocxLoaderImageExtraction:
    """Tests for inline image extraction and placeholder insertion."""

    def test_image_placeholder_inserted_in_text(self, tmp_path):
        """[IMAGE: xxx] placeholder is present in doc.text when image exists."""
        docx_path = _create_docx_with_image(tmp_path)
        loader = DocxLoader(
            extract_images=True,
            image_storage_dir=str(tmp_path / "images"),
        )
        doc = loader.load(docx_path)

        assert "[IMAGE:" in doc.text

    def test_image_metadata_recorded(self, tmp_path):
        """metadata['images'] is populated with at least one entry."""
        docx_path = _create_docx_with_image(tmp_path)
        loader = DocxLoader(
            extract_images=True,
            image_storage_dir=str(tmp_path / "images"),
        )
        doc = loader.load(docx_path)

        assert "images" in doc.metadata
        assert len(doc.metadata["images"]) >= 1

    def test_image_metadata_fields(self, tmp_path):
        """Each image metadata entry contains the required fields."""
        docx_path = _create_docx_with_image(tmp_path)
        loader = DocxLoader(
            extract_images=True,
            image_storage_dir=str(tmp_path / "images"),
        )
        doc = loader.load(docx_path)

        img = doc.metadata["images"][0]
        assert "id" in img
        assert "path" in img
        assert "text_offset" in img
        assert "text_length" in img

    def test_image_file_saved_to_disk(self, tmp_path):
        """Extracted image file is actually written to image_storage_dir."""
        docx_path = _create_docx_with_image(tmp_path)
        image_dir = tmp_path / "images"
        loader = DocxLoader(extract_images=True, image_storage_dir=str(image_dir))
        doc = loader.load(docx_path)

        img = doc.metadata["images"][0]
        assert Path(img["path"]).exists()

    def test_text_offset_points_to_placeholder(self, tmp_path):
        """text_offset in metadata correctly points to [IMAGE: ...] in text."""
        docx_path = _create_docx_with_image(tmp_path)
        loader = DocxLoader(
            extract_images=True,
            image_storage_dir=str(tmp_path / "images"),
        )
        doc = loader.load(docx_path)

        img = doc.metadata["images"][0]
        offset = img["text_offset"]
        length = img["text_length"]
        snippet = doc.text[offset : offset + length]
        assert snippet.startswith("[IMAGE:")

    def test_placeholder_between_surrounding_text(self, tmp_path):
        """Placeholder appears between 'before' and 'after' paragraphs."""
        docx_path = _create_docx_with_image(tmp_path)
        loader = DocxLoader(
            extract_images=True,
            image_storage_dir=str(tmp_path / "images"),
        )
        doc = loader.load(docx_path)

        before_pos = doc.text.find("Paragraph before image.")
        placeholder_pos = doc.text.find("[IMAGE:")
        after_pos = doc.text.find("Paragraph after image.")

        assert before_pos < placeholder_pos < after_pos

    def test_no_images_metadata_absent(self, tmp_path):
        """metadata['images'] key is absent when DOCX has no images."""
        docx_path = _create_test_docx(tmp_path)
        loader = DocxLoader(
            extract_images=True,
            image_storage_dir=str(tmp_path / "images"),
        )
        doc = loader.load(docx_path)

        assert "images" not in doc.metadata

    def test_extract_images_false_skips_images(self, tmp_path):
        """extract_images=False produces no placeholders or images metadata."""
        docx_path = _create_docx_with_image(tmp_path)
        loader = DocxLoader(
            extract_images=False,
            image_storage_dir=str(tmp_path / "images"),
        )
        doc = loader.load(docx_path)

        assert "[IMAGE:" not in doc.text
        assert "images" not in doc.metadata


class TestDocxLoaderHelperMethods:
    """Tests for helper methods."""

    def test_compute_file_hash_consistency(self, tmp_path):
        """_compute_file_hash returns consistent hash for same content."""
        docx_path = _create_test_docx(tmp_path)

        loader = DocxLoader()
        hash1 = loader._compute_file_hash(docx_path)
        hash2 = loader._compute_file_hash(docx_path)

        assert hash1 == hash2
        assert len(hash1) == 64

    def test_validate_file_via_subclass(self, tmp_path):
        """_validate_file works when accessed via DocxLoader."""
        docx_path = _create_test_docx(tmp_path)

        validated = DocxLoader._validate_file(docx_path)
        assert validated.exists()
        assert validated.is_file()
