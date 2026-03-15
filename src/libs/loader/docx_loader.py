"""Word (.docx) Loader implementation using python-docx.

This module implements DOCX parsing, extracting paragraphs, headings, and
inline images, converting them to standardized Markdown format.

Features:
- Heading detection (Heading 1-6) mapped to Markdown # ## ###
- Paragraph text extraction
- Inline image extraction: saves to disk, inserts [IMAGE: xxx] placeholders
- Metadata: source_path, doc_type, doc_hash, paragraph_count, images
"""

from __future__ import annotations

import hashlib
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

try:
    from docx import Document as DocxDocument  # type: ignore[reportMissingTypeStubs]
    from docx.oxml.ns import qn

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from src.core.types import Document
from src.libs.loader.base_loader import BaseLoader

logger = logging.getLogger(__name__)

_CONTENT_TYPE_EXT: Dict[str, str] = {
    "image/jpeg": "jpg",
    "image/png": "png",
    "image/gif": "gif",
    "image/bmp": "bmp",
    "image/tiff": "tiff",
    "image/x-wmf": "wmf",
    "image/x-emf": "emf",
    "image/svg+xml": "svg",
}


class DocxLoader(BaseLoader):
    """Word (.docx) Loader using python-docx.

    This loader:
    1. Extracts paragraphs and headings from DOCX
    2. Converts headings to Markdown (# ## ###)
    3. Extracts inline images, saves to disk, inserts [IMAGE: xxx] placeholders
    4. Returns Document with source_path, doc_type, doc_hash, paragraph_count,
       and optional images metadata
    """

    def __init__(
        self,
        extract_images: bool = True,
        image_storage_dir: str | Path = "data/images",
    ) -> None:
        """Initialize Docx Loader.

        Args:
            extract_images: Whether to extract and save inline images.
            image_storage_dir: Base directory for storing extracted images.
                               A subdirectory named after doc_hash is created
                               automatically under this path.
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DocxLoader. "
                "Install with: pip install python-docx"
            )
        self.extract_images = extract_images
        self.image_storage_dir = Path(image_storage_dir)

    def load(self, file_path: str | Path) -> Document:
        """Load and parse a DOCX file.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            Document with Markdown text and metadata.

        Raises:
            FileNotFoundError: If the DOCX file doesn't exist.
            ValueError: If the file is not a valid DOCX.
            RuntimeError: If parsing fails critically.
        """
        path = self._validate_file(file_path)
        if path.suffix.lower() != ".docx":
            raise ValueError(f"File is not a DOCX: {path}")

        doc_hash = self._compute_file_hash(path)
        doc_id = f"doc_{doc_hash[:16]}"

        try:
            doc = DocxDocument(str(path))
        except Exception as e:
            logger.error(f"Failed to parse DOCX {path}: {e}")
            raise RuntimeError(f"DOCX parsing failed: {e}") from e

        # Build rId -> image info mapping up front (graceful degradation on failure)
        rid_to_info: Dict[str, Dict[str, str]] = {}
        images_metadata: List[Dict[str, Any]] = []
        if self.extract_images:
            rid_to_info, images_metadata = self._extract_images(path, doc, doc_hash)

        # Iterate paragraphs: insert image placeholders before paragraph text
        parts: list[str] = []
        for para in doc.paragraphs:
            # Inline images appear as <w:drawing> inside paragraph runs;
            # detect them first so the placeholder preserves reading order.
            for rId in self._para_image_rids(para):
                if rId in rid_to_info:
                    parts.append(rid_to_info[rId]["placeholder"])

            text = para.text.strip()
            if not text:
                continue
            style = para.style.name or ""
            if style.startswith("Heading 1"):
                parts.append(f"# {text}")
            elif style.startswith("Heading 2"):
                parts.append(f"## {text}")
            elif style.startswith("Heading 3"):
                parts.append(f"### {text}")
            elif style.startswith("Heading 4"):
                parts.append(f"#### {text}")
            elif style.startswith("Heading 5"):
                parts.append(f"##### {text}")
            elif style.startswith("Heading 6"):
                parts.append(f"###### {text}")
            else:
                parts.append(text)

        text_content = "\n\n".join(parts) if parts else ""

        # Back-fill accurate text_offset now that the final text is assembled
        for img_meta in images_metadata:
            placeholder = f"[IMAGE: {img_meta['id']}]"
            offset = text_content.find(placeholder)
            if offset >= 0:
                img_meta["text_offset"] = offset

        title = self._extract_title(text_content)
        metadata: Dict[str, Any] = {
            "source_path": str(path),
            "doc_type": "docx",
            "doc_hash": doc_hash,
            "paragraph_count": len(doc.paragraphs),
        }
        if title:
            metadata["title"] = title
        if images_metadata:
            metadata["images"] = images_metadata

        return Document(
            id=doc_id,
            text=text_content,
            metadata=metadata,
        )

    # ------------------------------------------------------------------
    # Image helpers
    # ------------------------------------------------------------------

    def _extract_images(
        self,
        path: Path,
        doc: "DocxDocument",
        doc_hash: str,
    ) -> Tuple[Dict[str, Dict[str, str]], List[Dict[str, Any]]]:
        """Extract all inline images from the DOCX and save to disk.

        Returns:
            (rid_to_info, images_metadata)
            - rid_to_info: maps each relationship ID to
              {"image_id", "placeholder", "path"}
            - images_metadata: list of image metadata dicts compatible with
              the Document.metadata["images"] schema used by PdfLoader
        """
        rid_to_info: Dict[str, Dict[str, str]] = {}
        images_metadata: List[Dict[str, Any]] = []

        try:
            image_dir = self.image_storage_dir / doc_hash
            image_dir.mkdir(parents=True, exist_ok=True)

            seq = 0
            for rel in doc.part.rels.values():
                if "image" not in rel.reltype:
                    continue

                rId = rel.rId
                image_part = rel.target_part
                image_bytes: bytes = image_part.blob
                ext = _CONTENT_TYPE_EXT.get(image_part.content_type, "png")

                seq += 1
                image_id = f"{doc_hash[:8]}_0_{seq}"
                image_path = image_dir / f"{image_id}.{ext}"

                with open(image_path, "wb") as f:
                    f.write(image_bytes)

                try:
                    rel_path: Path = image_path.relative_to(Path.cwd())
                except ValueError:
                    rel_path = image_path.absolute()

                placeholder = f"[IMAGE: {image_id}]"
                rid_to_info[rId] = {
                    "image_id": image_id,
                    "placeholder": placeholder,
                    "path": str(rel_path),
                }
                images_metadata.append(
                    {
                        "id": image_id,
                        "path": str(rel_path),
                        "page": None,
                        "text_offset": 0,  # updated after text assembly
                        "text_length": len(placeholder),
                        "position": {"sequence": seq},
                    }
                )

            if images_metadata:
                logger.info(
                    f"Extracted {len(images_metadata)} image(s) from {path}"
                )

        except Exception as e:
            logger.warning(
                f"Image extraction failed for {path}, continuing text-only: {e}"
            )
            return {}, []

        return rid_to_info, images_metadata

    @staticmethod
    def _para_image_rids(para: Any) -> List[str]:
        """Return the r:embed rIds of inline images inside a paragraph."""
        rids: List[str] = []
        for blip in para._element.iter(qn("a:blip")):
            rId = blip.get(qn("r:embed"))
            if rId:
                rids.append(rId)
        return rids

    # ------------------------------------------------------------------
    # Shared helpers
    # ------------------------------------------------------------------

    def _compute_file_hash(self, file_path: Path) -> str:
        """Compute SHA256 hash of file content."""
        sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                sha256.update(chunk)
        return sha256.hexdigest()

    def _extract_title(self, text: str) -> Optional[str]:
        """Extract title from first Markdown heading or first non-empty line."""
        lines = text.split("\n")
        for line in lines[:20]:
            line = line.strip()
            if line.startswith("# "):
                return line[2:].strip()
        for line in lines[:10]:
            line = line.strip()
            if line:
                return line
        return None
